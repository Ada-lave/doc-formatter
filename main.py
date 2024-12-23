from fastapi import FastAPI, File, UploadFile, Query
from fastapi.responses import JSONResponse
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware
from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
import tempfile
import os
import subprocess

app = FastAPI()

# Настройка CORS (опционально)
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


def insert_paragraph_after(paragraph, text=None, style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    if style is not None:
        new_para.style = style


def is_paragraph_picture(paragraph: Paragraph):
    """
    Проверяет, содержит ли абзац картинку.
    """
    return any("pic:pic" in run.element.xml for run in paragraph.runs)


@app.post("/process-docx")
async def process_docx(
    file: UploadFile = File(...),
    top_margin: bool = Query(default=False),
    bottom_margin: bool = Query(default=True),
    convert_to_pdf: bool = Query(default=False),
):
    if not file.filename.endswith(".docx"):
        return JSONResponse({"error": "Uploaded file is not a .docx file"}, 422)

    # Создаём временный файл для сохранения загруженного содержимого
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_file:
        temp_file.write(await file.read())
        temp_file_path = temp_file.name

    # Открываем документ и проверяем содержимое
    doc = Document(temp_file_path)
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.lower().startswith("рис"):
            next_paragraph_index = i + 1
            next_paragraph = doc.paragraphs[next_paragraph_index]
            # Если после подрисуночной надписи нет отступа, добавляем
            if next_paragraph.text.strip() != "" or is_paragraph_picture(
                next_paragraph
            ):
                if bottom_margin:
                    insert_paragraph_after(paragraph, "")
            # Проверяем наличие отступа сверху от картинки
            if top_margin and i > 2:
                prev_paragraph = doc.paragraphs[i - 2]
                if len(prev_paragraph.text.strip()) > 0 and not is_paragraph_picture(
                    prev_paragraph
                ):
                    insert_paragraph_after(prev_paragraph, "")

    output_file_path = tempfile.mktemp(suffix=".docx")
    doc.save(output_file_path)

    if convert_to_pdf:
        pdf_file_path = os.path.splitext(output_file_path)[0] + ".pdf"
        subprocess.run(["libreoffice", "--headless", "--convert-to", "pdf", output_file_path, "--outdir", os.path.dirname(output_file_path)])
        os.remove(output_file_path)
        return FileResponse(
            pdf_file_path,
            filename=f"processed_{os.path.splitext(file.filename)[0]}.pdf",
            media_type="application/pdf",
        )

    os.remove(temp_file_path)

    return FileResponse(
        output_file_path,
        filename=f"processed_{file.filename}",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


# Запуск приложения (только при локальном запуске)
if __name__ == "__main__":
    import uvicorn

    uvicorn.run(app, host="0.0.0.0", port=8000)